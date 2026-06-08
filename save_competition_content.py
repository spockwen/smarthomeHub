# -*- coding: utf-8 -*-
"""
读取 docx 文件的完整内容（包括段落和表格）并保存到文件
"""
from docx import Document

doc = Document(r'E:\harmonyos\A9-基于OpenHarmony的家居设备控制系统.docx')

output_path = r'E:\harmonyos\SmartHomeAIHub\competition_full_content.txt'

with open(output_path, 'w', encoding='utf-8') as f:
    f.write("=" * 60 + "\n")
    f.write("段落内容:\n")
    f.write("=" * 60 + "\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"[{i}] {para.text}\n")

    f.write("\n" + "=" * 60 + "\n")
    f.write("表格内容:\n")
    f.write("=" * 60 + "\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\n--- 表格 {t_idx + 1} ---\n")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(row_data) + "\n")

print(f"内容已保存到: {output_path}")