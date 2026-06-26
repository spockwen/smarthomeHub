# -*- coding: utf-8 -*-
"""
读取 docx 文件的完整内容（包括段落和表格）
"""
from docx import Document

doc = Document(r'E:\harmonyos\A9-基于OpenHarmony的家居设备控制系统.docx')

print("=" * 60)
print("段落内容:")
print("=" * 60)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print("\n" + "=" * 60)
print("表格内容:")
print("=" * 60)
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx + 1} ---")
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        print(" | ".join(row_data))